"""
Text extraction module for PDF and DOCX files.

Handles parsing of resume files and extracting clean text.
Fix: clean_text() now preserves newlines so section detection works correctly.
"""

import re
import logging
import tempfile
import os
from typing import Optional
from pathlib import Path

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

MAX_TEXT_LENGTH = 50_000  # safety cap: ~50k chars is more than any real CV


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """
    Extract text from a PDF file using pdfplumber.

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted and cleaned text, or None if extraction fails

    Raises:
        FileNotFoundError: If file does not exist
        ValueError: If pdfplumber is not installed
    """
    if pdfplumber is None:
        raise ValueError(
            "pdfplumber is not installed. Run: pip install pdfplumber"
        )

    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"PDF file not found: {file_path}")

    try:
        pages_text = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text)

        extracted_text = "\n\n".join(pages_text)
        cleaned = clean_text(extracted_text)
        logger.info(f"Extracted text from PDF: {file_path} ({len(cleaned)} chars)")
        return cleaned[:MAX_TEXT_LENGTH]

    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}")
        raise


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    Extract text from a DOCX file using python-docx.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted and cleaned text, or None if extraction fails

    Raises:
        FileNotFoundError: If file does not exist
        ValueError: If python-docx is not installed
    """
    if Document is None:
        raise ValueError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    try:
        doc = Document(file_path)
        lines = []

        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())

        extracted_text = "\n".join(lines)
        cleaned = clean_text(extracted_text)
        logger.info(f"Extracted text from DOCX: {file_path} ({len(cleaned)} chars)")
        return cleaned[:MAX_TEXT_LENGTH]

    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}")
        raise


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> Optional[str]:
    """
    Extract text directly from file bytes (safe temp-file approach).

    Uses a proper temp file with guaranteed cleanup via context manager.

    Args:
        file_bytes: Raw bytes of the uploaded file
        filename: Original filename (used to detect extension)

    Returns:
        Extracted text string
    """
    ext = Path(filename).suffix.lower()
    suffix = ext if ext in (".pdf", ".docx") else ".tmp"

    # NamedTemporaryFile with delete=False so we control deletion
    tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
    try:
        tmp.write(file_bytes)
        tmp.flush()
        tmp.close()

        if suffix == ".pdf":
            return extract_text_from_pdf(tmp.name)
        elif suffix == ".docx":
            return extract_text_from_docx(tmp.name)
        else:
            raise ValueError(f"Unsupported file type: {filename}")

    finally:
        # Always clean up, even if extraction raises
        try:
            os.unlink(tmp.name)
        except OSError:
            pass


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text.

    Fix: preserves single newlines so section headers can be found by
    extract_sections(). Previously collapsed ALL whitespace to a single
    space, which destroyed section structure.

    Args:
        text: Raw extracted text

    Returns:
        Cleaned text with preserved line breaks
    """
    # Remove Windows-style line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Remove control characters (keep \n = 0x0A, \t = 0x09)
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", text)

    # Collapse runs of spaces/tabs on the same line (but NOT newlines)
    text = re.sub(r"[ \t]+", " ", text)

    # Collapse more than 2 consecutive newlines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.splitlines()]
    text = "\n".join(lines)

    return text.strip()


def extract_sections(text: str) -> dict:
    """
    Extract common resume sections.

    Args:
        text: Full CV text (must retain newlines — use clean_text output)

    Returns:
        Dictionary with extracted sections
    """
    sections = {
        "summary": "",
        "experience": "",
        "education": "",
        "skills": "",
        "projects": "",
        "certifications": "",
        "full_text": text,
    }

    section_keywords = {
        "summary": ["summary", "objective", "profile", "about"],
        "experience": ["experience", "employment", "work history", "work experience"],
        "education": ["education", "academic background", "qualifications"],
        "skills": ["skills", "technical skills", "competencies", "technologies"],
        "projects": ["projects", "portfolio", "personal projects"],
        "certifications": ["certifications", "certificates", "awards", "achievements"],
    }

    all_keywords = [kw for kws in section_keywords.values() for kw in kws]
    boundary = "|".join(re.escape(k) for k in all_keywords)
    boundary_pattern = rf"(?:^|\n)(?:{boundary})\s*[:\-]?\s*(?:\n|$)"

    for section_name, keywords in section_keywords.items():
        for keyword in keywords:
            start_pattern = rf"(?:^|\n){re.escape(keyword)}\s*[:\-]?\s*(?:\n|$)"
            match = re.search(start_pattern, text, re.IGNORECASE)
            if match:
                start = match.end()
                # Find where next section starts
                next_match = re.search(boundary_pattern, text[start:], re.IGNORECASE)
                end = start + next_match.start() if next_match else len(text)
                sections[section_name] = text[start:end].strip()
                break

    return sections


def get_text_statistics(text: str) -> dict:
    """
    Calculate statistics about the text.

    Args:
        text: Input text

    Returns:
        Dictionary with statistics
    """
    words = text.split()
    sentences = re.split(r"[.!?]+", text)

    return {
        "word_count": len(words),
        "character_count": len(text),
        "sentence_count": len([s for s in sentences if s.strip()]),
        "average_word_length": round(len(text) / len(words), 2) if words else 0,
    }